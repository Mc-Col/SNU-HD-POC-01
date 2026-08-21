# -*- coding: utf-8 -*-
"""
docx 표 셀 주입 엔진 — 원본 서식(글꼴/크기/정렬) 보존

셀 텍스트 규칙
  '\n'   : 줄바꿈 (간격 없이 붙임)
  '\n\n' : 논리 블록 구분 (다음 줄 위에 여백 4pt)
"""
import copy
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run

BLOCK_GAP = Pt(4)   # 블록 사이 여백
LINE_GAP  = Pt(0)   # 줄 사이 여백


def _strip_runs(el):
    for tag in ("w:r", "w:hyperlink"):
        for r in el.findall(qn(tag)):
            el.remove(r)
    return el


def _run_from_tmpl(p, tmpl, text):
    if tmpl is None:
        p.add_run(text)
        return
    newr = copy.deepcopy(tmpl)
    for tag in ("w:t", "w:br", "w:tab", "w:cr"):
        for n in newr.findall(qn(tag)):
            newr.remove(n)
    p._p.append(newr)
    Run(newr, p).text = text


def set_cell(cell, text):
    """셀 내용을 text로 교체. 빈 줄은 블록 구분 여백으로 처리."""
    text = "" if text is None else str(text)
    lines = text.split("\n")

    paras = cell.paragraphs
    para_tmpl = copy.deepcopy(paras[0]._p)
    run_tmpl = None
    for p in paras:
        if p.runs:
            run_tmpl = copy.deepcopy(p.runs[0]._element)
            break
    for p in paras[1:]:
        p._p.getparent().remove(p._p)

    base = cell.paragraphs[0]
    _strip_runs(base._p)

    prev_el, first, gap = None, True, False
    made = 0
    for ln in lines:
        if ln == "":                 # 빈 줄 = 블록 구분 신호
            gap = True
            continue
        if first:
            pp, first = base, False
        else:
            newp = _strip_runs(copy.deepcopy(para_tmpl))
            prev_el.addnext(newp)
            pp = Paragraph(newp, base._parent)
        _run_from_tmpl(pp, run_tmpl, ln)
        fmt = pp.paragraph_format
        fmt.space_after = LINE_GAP
        fmt.space_before = BLOCK_GAP if gap and made else LINE_GAP
        gap = False
        prev_el = pp._p
        made += 1

    if made == 0:
        _run_from_tmpl(base, run_tmpl, "")


def add_rows(table, n):
    """마지막 행 서식을 복제해 n행 추가 후 내용 비움."""
    for _ in range(n):
        src = table.rows[-1]._tr
        src.addnext(copy.deepcopy(src))
        for c in table.rows[-1].cells:
            set_cell(c, "")


def fill(table, rows, start=1):
    """rows: [[cell,...],...] 를 start행부터 채움. 부족하면 행 추가. None = 원본 유지."""
    need = start + len(rows) - len(table.rows)
    if need > 0:
        add_rows(table, need)
    for i, r in enumerate(rows):
        tr = table.rows[start + i]
        for j, v in enumerate(r):
            if v is not None and j < len(tr.cells):
                set_cell(tr.cells[j], v)
