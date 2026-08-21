# -*- coding: utf-8 -*-
"""
워크북 빌드 스크립트

  단일 소스(content_*.py) → 두 산출물
    1) (integrated)AI_AX_problem_formulation_workbook_rev.X.docx  (제출용)
    2) docs/workbook_rev.X.md                                     (리뷰·diff용)

  사용법:
    python tools/build.py            # 현재 REV로 빌드
    python tools/build.py rev.B      # 리비전 지정 빌드

  리비전 갱신 시 docs/change_log.md 에 변경 이력을 함께 기록할 것.
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from engine import fill
import content_ac as AC
import content_df as DF
import content_gk as GK

ROOT    = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ARGS    = [a for a in sys.argv[1:] if not a.startswith("--")]
COMPACT = "--compact" in sys.argv
REV     = ARGS[0] if ARGS else "rev.B"
SUFFIX  = "_compact" if COMPACT else ""
SRC     = os.path.join(ROOT, "data", "(integrated)AI_AX_problem_formulation_workbook_blank.docx")
OUT     = os.path.join(ROOT, f"(integrated)AI_AX_problem_formulation_workbook_{REV}{SUFFIX}.docx")
MD      = os.path.join(ROOT, "docs", f"workbook_{REV}.md")

# 표 인덱스 → 내용. 주입하지 않는 표(안내문·작성틀)는 목록에서 제외
MAP = [
    (0,  AC.COVER, "표지"),
    (3,  AC.A1,    "A-1 문제 기본정보"),
    (4,  AC.A2,    "A-2 현재 문제 상황"),
    (6,  AC.A3,    "A-3 문제 명세 초안"),
    (7,  AC.B,     "Part B 의사결정 구조"),
    (8,  AC.B1,    "B-1 개선 대상"),
    (9,  AC.C,     "Part C 적용 여건"),
    (10, AC.C1,    "C-1 종합 검토"),
    (11, DF.D,     "Part D 요구 기능"),
    (12, DF.D1,    "D-1 기능 우선순위"),
    (13, DF.E1,    "E-1 필요 데이터"),
    (14, DF.E2,    "E-2 데이터 확보 상태"),
    (15, DF.E3,    "E-3 부족 시 대안"),
    (16, DF.F,     "Part F 방법론 비교"),
    (17, DF.F1,    "F-1 접근법 선정"),
    (18, GK.G,     "Part G 출력·운영 통합"),
    (19, GK.G1,    "G-1 입력-처리-출력-반영"),
    (20, GK.H,     "Part H KPI"),
    (21, GK.H1,    "H-1 검증 설계"),
    (22, GK.I,     "Part I 리스크"),
    (23, GK.I1,    "I-1 핵심 선행조건"),
    (24, GK.J,     "Part J PoC 범위"),
    (25, GK.J1,    "J-1 실행 로드맵"),
    (26, GK.K1,    "K-1 연구개발 논리"),
    (28, GK.K2,    "K-2 통합 문제 정식화"),
    (29, GK.K3,    "K-3 과제 개요"),
]


def allow_row_split(doc):
    """모든 행의 cantSplit(행 분할 금지)을 제거.

    원본 양식은 164개 행 전부에 cantSplit이 걸려 있어, 행 높이가 남은 페이지
    공간보다 크면 행 전체가 다음 페이지로 밀리고 그 위에 빈 공간이 남는다.
    내용이 긴 셀에서는 이 낭비가 매우 커지므로 분할을 허용한다.
    """
    from docx.oxml.ns import qn
    n = 0
    for t in doc.tables:
        for r in t.rows:
            trPr = r._tr.find(qn("w:trPr"))
            if trPr is None:
                continue
            for cs in trPr.findall(qn("w:cantSplit")):
                trPr.remove(cs)
                n += 1
    return n


def drop_page_breaks(doc):
    """Part 시작 전 페이지 나누기를 제거해 내용이 연속 흐르게 함.

    원본 양식은 Part마다 새 페이지에서 시작하도록 되어 있어, 표가 조금만
    넘쳐도 그 아래가 통째로 비는 페이지가 생긴다(측정 결과 10페이지).
    분량을 우선할 때만 사용.
    """
    from docx.oxml.ns import qn
    body = doc.element.body
    n = 0
    for br in body.findall(".//" + qn("w:br")):
        if br.get(qn("w:type")) == "page":
            br.getparent().remove(br); n += 1
    for pbb in body.findall(".//" + qn("w:pageBreakBefore")):
        pbb.getparent().remove(pbb); n += 1
    return n


def emit_md(doc, rev, path):
    """빌드된 docx를 순서대로 읽어 마크다운으로 출력 (docx와 항상 일치)"""
    body = doc.element.body
    out = [f"# HD AI/AX 현업 문제 발굴·정의 워크북 — {rev}", "",
           "> 이 파일은 `tools/build.py`가 docx에서 자동 생성함. 직접 편집하지 말고",
           "> `tools/content_*.py`를 수정한 뒤 재빌드할 것.", ""]
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            t = Paragraph(child, doc).text.strip()
            if t:
                out.append(t if len(t) > 40 else f"### {t}")
                out.append("")
        elif tag == "tbl":
            tb = Table(child, doc)
            rows = [[c.text.strip().replace("\n", "<br>") or "&nbsp;" for c in r.cells]
                    for r in tb.rows]
            if not rows:
                continue
            ncol = max(len(r) for r in rows)
            for r in rows:
                r += ["&nbsp;"] * (ncol - len(r))
            out.append("| " + " | ".join(rows[0]) + " |")
            out.append("|" + "---|" * ncol)
            for r in rows[1:]:
                out.append("| " + " | ".join(r) + " |")
            out.append("")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(out))
    return len(out)


def main():
    d = Document(SRC)
    print(f"리비전: {REV}")
    print(f"{'표':<6}{'항목':<26}{'기존행':>6}{'주입':>6}{'추가':>6}")
    print("-" * 52)
    added_total = 0
    for idx, rows, label in MAP:
        t = d.tables[idx]
        before = len(t.rows)
        fill(t, rows, start=1)
        added = len(t.rows) - before
        added_total += added
        print(f"T{idx:<5}{label:<26}{before:>6}{len(rows):>6}{added:>6}")
    split = allow_row_split(d)
    print("-" * 52)
    print(f"행 분할 허용(cantSplit 제거): {split}개 행")
    if COMPACT:
        print(f"페이지 나누기 제거(연속 흐름): {drop_page_breaks(d)}개")
    try:
        d.save(OUT)
    except PermissionError:
        sys.exit(f"\n[중단] 저장 실패 — 파일이 잠겨 있음:\n  {OUT}\n"
                 f"  Word/한글에서 해당 문서를 닫은 뒤 다시 실행할 것.")
    print("-" * 52)
    print(f"표 {len(MAP)}개 주입, 행 {added_total}개 추가")

    chars = sum(len(c.text.strip())
                for i, t in enumerate(d.tables) if i in {m[0] for m in MAP}
                for r in t.rows for c in r.cells)
    empties = [(i, ri, ci)
               for i, t in enumerate(d.tables) if i in {m[0] for m in MAP}
               for ri, r in enumerate(t.rows) if ri > 0
               for ci, c in enumerate(r.cells) if not c.text.strip()]
    lines = emit_md(Document(OUT), REV, MD)

    print(f"\ndocx : {OUT}  ({os.path.getsize(OUT):,} bytes)")
    print(f"md   : {MD}  ({lines} lines)")
    print(f"본문 : {chars:,}자 / 빈 셀 {len(empties)}개 {empties}")


if __name__ == "__main__":
    main()
