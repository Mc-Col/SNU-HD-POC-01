# -*- coding: utf-8 -*-
"""
과제제안서 작성 프롬프트 로그 (제출용) 생성

  python tools/build_promptlog.py

  Readme 원문은 Readme/*.pdf 에서 직접 추출하므로 전사 오류가 없고 재생성 가능함.
"""
import os, sys, glob
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT  = os.path.join(ROOT, "과제제안서_프롬프트로그_rev.A.docx")

KO      = "맑은 고딕"
NAVY    = RGBColor(0x1F, 0x30, 0x64)
GREY    = RGBColor(0x55, 0x55, 0x55)
SHADE   = "F2F4F8"   # 프롬프트 원문 박스
SHADE2  = "FFF8E7"   # 응답 핵심 박스


# ────────────────────────────── 서식 헬퍼 ──────────────────────────────
def set_font(run, size=10, bold=False, color=None, name=KO):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), name)


def shade(par, hexcolor):
    pPr = par._p.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hexcolor)
    pPr.append(sh)


def border(par, sides=("top", "left", "bottom", "right"), hexcolor="C8CEDC"):
    """단락 테두리. 여러 단락을 하나의 박스로 보이게 하려면
    첫 줄 top+left+right / 중간 left+right / 마지막 bottom+left+right 로 지정."""
    pPr = par._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        if side in sides:
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
            e.set(qn("w:space"), "4"); e.set(qn("w:color"), hexcolor)
        else:
            e.set(qn("w:val"), "nil")
        bd.append(e)
    pPr.append(bd)


def P(doc, text="", size=10, bold=False, color=None, before=0, after=3,
      indent=0, align=None, line=1.35):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing = line
    if indent: pf.left_indent = Cm(indent)
    if align is not None: p.alignment = align
    if text:
        set_font(p.add_run(text), size, bold, color)
    return p


def H(doc, text, level=1):
    spec = {1: (15, 14, 6), 2: (12, 12, 4), 3: (10.5, 9, 3)}[level]
    p = P(doc, text, size=spec[0], bold=True, color=NAVY,
          before=spec[1], after=spec[2], line=1.2)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        bd = OxmlElement("w:pBdr"); e = OxmlElement("w:bottom")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "8")
        e.set(qn("w:space"), "3"); e.set(qn("w:color"), "1F3064")
        bd.append(e); pPr.append(bd)
    return p


def BOX(doc, lines, fill=SHADE, size=9.5, mono=False):
    """여러 단락을 하나의 음영 박스로 렌더링. lines: 문자열 리스트"""
    name = "Consolas" if mono else KO
    last = len(lines) - 1
    for i, ln in enumerate(lines):
        p = P(doc, "", size=size, before=0, after=0, indent=0.2, line=1.3)
        if ln:
            set_font(p.add_run(ln), size, False, None, name)
        shade(p, fill)
        sides = ["left", "right"]
        if i == 0: sides.append("top")
        if i == last: sides.append("bottom")
        border(p, tuple(sides))
    P(doc, "", after=4)


def BUL(doc, items, size=10, indent=0.5):
    for it in items:
        P(doc, "· " + it, size=size, indent=indent, after=2)


def TBL(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        set_font(c.paragraphs[0].add_run(h), 9.5, True, RGBColor(0xFF, 0xFF, 0xFF))
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
        sh.set(qn("w:fill"), "1F3064"); c._tc.get_or_add_tcPr().append(sh)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = ""
            for j, ln in enumerate(str(v).split("\n")):
                p = cells[i].paragraphs[0] if j == 0 else cells[i].add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.25
                set_font(p.add_run(ln), 9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    P(doc, "", after=6)
    return t


def readme_text():
    pdf = glob.glob(os.path.join(ROOT, "Readme", "*.pdf"))
    if not pdf:
        return ["(Readme PDF를 찾을 수 없음)"]
    import pymupdf
    d = pymupdf.open(pdf[0])
    out = []
    for pg in d:
        for ln in pg.get_text().split("\n"):
            ln = ln.rstrip()
            if ln.strip():
                out.append(ln)
    return out


# ────────────────────────────── 문서 생성 ──────────────────────────────
def main():
    import promptlog_content as C

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.2)
    st = doc.styles["Normal"]
    st.font.name = KO; st.font.size = Pt(10)
    rf = st.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), KO)

    # 표지
    P(doc, "HD-SNU AI 고급과정 · 텀 프로젝트", size=10, color=GREY, after=2)
    P(doc, "과제제안서 작성 프롬프트 로그", size=22, bold=True, color=NAVY,
      before=6, after=4, line=1.15)
    P(doc, C.TITLE, size=11.5, bold=True, after=14, line=1.3)
    TBL(doc, ["구분", "내용"],
        [["과제명", C.TITLE],
         ["산출물", "(integrated)AI_AX_problem_formulation_workbook rev.C — Part A~K"],
         ["작성 도구", "Claude Code (Anthropic) / 초기 컨텍스트 프롬프트는 Gemini와 사전 작성"],
         ["작성 기간", "2026-08-20 (1일, 단계별 검토 반복)"],
         ["프롬프트 수", f"P0(초기 컨텍스트) + P1~P{len(C.LOG)} (단계별 {len(C.LOG)}회)"],
         ["작성자", ""],
         ["소속/부서", ""]],
        widths=[3.2, 13.0])

    H(doc, "0. 문서 개요", 1)
    P(doc, "본 문서는 과제제안서(AI/AX 현업 문제 발굴·정의 워크북) 작성에 사용한 프롬프트와 "
           "그에 대한 응답의 핵심, 그리고 각 단계가 워크북에 반영된 위치를 기록한 것임. "
           "일괄 지시가 아니라 단계별 검토·교정 방식으로 작성했으며, 모델 제안을 도메인 지식으로 "
           "교정한 지점을 함께 남겼음.")
    H(doc, "프롬프트 운용 원칙", 3)
    BUL(doc, C.PRINCIPLES)

    # P0 Readme
    doc.add_page_break()
    H(doc, "1. 초기 컨텍스트 프롬프트 (P0)", 1)
    P(doc, "작성 착수 전 도메인 맥락·문제 정의·솔루션 방향을 정리한 프롬프트임. Gemini와 사전 작성하여 "
           "Claude Code에 최초 입력으로 전달했음. 원문 전문을 아래에 첨부함.", after=6)
    P(doc, "※ 이 프롬프트가 기술을 제약으로 확정한 부분은 P1의 검토에서 가설로 강등되었음 "
           "(4장 회고 ① 참조).", size=9.5, color=GREY, after=8)
    BOX(doc, readme_text(), fill=SHADE, size=9, mono=False)

    # 프롬프트 로그
    doc.add_page_break()
    H(doc, "2. 프롬프트 로그", 1)
    for i, (title, prompt, resp, applied) in enumerate(C.LOG):
        if i: P(doc, "", after=8)
        H(doc, title, 2)
        P(doc, "프롬프트", size=9.5, bold=True, color=NAVY, before=2, after=2)
        BOX(doc, prompt, fill=SHADE, size=9.5)
        P(doc, "응답 핵심", size=9.5, bold=True, color=NAVY, before=6, after=2)
        for r in resp:
            ind = 0.75 if r.startswith(("·", "①", "②", "③", "④", "⑤", "    ")) else 0.35
            P(doc, r, size=9.5, indent=ind, after=2, line=1.3)
        P(doc, "반영 위치", size=9.5, bold=True, color=NAVY, before=6, after=2)
        p = P(doc, applied, size=9.5, indent=0.35, after=4)
        shade(p, SHADE2)

    # 회고
    doc.add_page_break()
    H(doc, "3. 프롬프트 관리 회고", 1)
    for t, lines in C.RETRO:
        H(doc, t, 2)
        for ln in lines:
            P(doc, ln, size=9.5, indent=(0.75 if ln.startswith("·") else 0.35),
              after=3, line=1.35)

    # 산출물
    H(doc, "4. 산출물 목록", 1)
    TBL(doc, ["구분", "파일", "내용"], C.OUTPUTS, widths=[3.4, 6.6, 6.2])
    P(doc, "워크북 내용은 tools/content_*.py를 단일 소스로 하여 build.py가 docx와 markdown을 "
           "생성하는 구조임. 재빌드 시 산출물이 바이트 단위로 동일함을 확인했음.",
      size=9.5, color=GREY)

    doc.save(OUT)
    print(f"저장: {OUT}")
    print(f"크기: {os.path.getsize(OUT):,} bytes")


if __name__ == "__main__":
    main()
